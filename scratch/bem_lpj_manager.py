import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys
import re
import json
import zipfile
import tempfile

def clean_currency(val):
    if val is None:
        return 0
    if isinstance(val, (int, float)):
        return int(val)
    val_clean = str(val).replace("Rp", "").replace(".", "").replace(",", "").strip()
    if not val_clean:
        return 0
    try:
        return int(float(val_clean))
    except ValueError:
        return 0

def format_currency(val):
    return f"Rp {val:,}".replace(",", ".")

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def clean_proker_name(name):
    if not name:
        return ""
    return re.sub(r'^\d+[\.\s-]*', '', name).strip()

def parse_id_date(d):
    if not d:
        return 0
    match_iso = re.search(r'^(\d{4})-(\d{2})-(\d{2})$', d.strip())
    if match_iso:
        return int(match_iso.group(1)) * 10000 + int(match_iso.group(2)) * 100 + int(match_iso.group(3))
        
    id_months = {
        'januari': 1, 'februari': 2, 'maret': 3, 'april': 4, 'mei': 5, 'juni': 6,
        'juli': 7, 'agustus': 8, 'september': 9, 'oktober': 10, 'november': 11, 'desember': 12
    }
    match = re.search(r'(\d+)(?:\s*[-–]\s*\d+)?\s+(\w+)\s+(\d{4})', d.strip().lower())
    if match:
        day = int(match.group(1))
        month_name = match.group(2)
        year = int(match.group(3))
        month = id_months.get(month_name, 1)
        return year * 10000 + month * 100 + day
    return 0

def format_id_date_str(d):
    if not d:
        return ""
    match_iso = re.search(r'^(\d{4})-(\d{2})-(\d{2})$', d.strip())
    if match_iso:
        year = match_iso.group(1)
        month = int(match_iso.group(2))
        day = int(match_iso.group(3))
        months = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
        return f"{day} {months[month]} {year}"
    return d

def format_kementerian_title(k_name):
    if not k_name:
        return "MENTERI KEMENTERIAN"
    k_name = k_name.upper().strip()
    if k_name.startswith("MENTERI MENTERI"):
        k_name = k_name[15:].strip()
    if k_name.startswith("MENTERI "):
        k_name = k_name[8:].strip()
    if k_name.startswith("KEMENTERIAN "):
        k_name = k_name[12:].strip()
    return f"MENTERI {k_name}"

def extract_periode_years(periode_str):
    if not periode_str:
        return ""
    match = re.search(r'\d{4}-\d{4}', periode_str)
    if match:
        return match.group(0)
    return periode_str

def resolve_photo_path(photo_path):
    if not photo_path:
        return ""
    # If it's already an absolute URL (S3), download it to a temp file
    if photo_path.startswith("http://") or photo_path.startswith("https://"):
        return _download_to_temp(photo_path)
    if os.path.exists(photo_path) and os.path.isfile(photo_path):
        return photo_path
    uploads_pos = photo_path.find('uploads/')
    if uploads_pos != -1:
        rel_path = photo_path[uploads_pos:]
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        local_path = os.path.join(project_root, rel_path)
        if os.path.exists(local_path) and os.path.isfile(local_path):
            return local_path
    # Fallback: try S3 public URL if configured
    s3_public_url = os.environ.get("S3_PUBLIC_URL", "")
    storage_method = os.environ.get("STORAGE_METHOD", "local")
    if storage_method == "s3" and s3_public_url:
        # Strip 'uploads/' prefix if present
        clean_key = photo_path
        if 'uploads/' in clean_key:
            clean_key = clean_key[clean_key.find('uploads/') + 8:]
        clean_key = clean_key.lstrip('/')
        full_url = s3_public_url.rstrip('/') + '/' + clean_key
        return _download_to_temp(full_url)
    return ""

def _download_to_temp(url):
    """Download a remote image to a temporary file and return the path."""
    try:
        import urllib.request
        ext = os.path.splitext(url.split('?')[0])[1] or '.png'
        fd, temp_path = tempfile.mkstemp(suffix=ext)
        os.close(fd)
        req = urllib.request.Request(
            url, 
            headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
            }
        )
        import ssl
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        
        with urllib.request.urlopen(req, context=ctx) as response, open(temp_path, 'wb') as out_file:
            shutil.copyfileobj(response, out_file)
        if os.path.getsize(temp_path) > 0:
            return temp_path
        os.remove(temp_path)
        return ""
    except Exception as e:
        print(f"Warning: Failed to download image from {url}: {e}")
        return ""

def get_docx_compatible_image(photo_path):
    if not photo_path or not os.path.exists(photo_path):
        return None, False
    ext = os.path.splitext(photo_path)[1].lower()
    if ext == '.webp':
        try:
            from PIL import Image
            img = Image.open(photo_path)
            fd, temp_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGBA')
            else:
                img = img.convert('RGB')
            img.save(temp_path, 'PNG')
            return temp_path, True
        except Exception as e:
            print(f"Error converting WebP image: {e}")
            return None, False
    return photo_path, False

def int_to_roman(num):
    val = [10, 9, 5, 4, 1]
    syb = ["X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

def get_bem_logo_path():
    # Try different possible paths relative to the script location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(script_dir, '..', 'assets', 'images', 'favicon', 'apple-touch-icon.png'),
        os.path.join(script_dir, '..', 'assets', 'images', 'favicon', 'web-app-manifest-512x512.png'),
        os.path.join(script_dir, '..', 'assets', 'images', 'favicon', 'web-app-manifest-192x192.png'),
        os.path.join(script_dir, '..', 'assets', 'images', 'favicon', 'favicon-96x96.png')
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None

def add_minister_cover(doc, kementerian_name, periode_years):
    # Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.15
    p_title.paragraph_format.space_before = Pt(72)  # Generous top margin
    p_title.paragraph_format.space_after = Pt(6)
    
    run_t1 = p_title.add_run("LAPORAN PERTANGGUNGJAWABAN\n")
    format_run(run_t1, font_name="Times New Roman", size_pt=14, bold=True)
    
    formatted_menteri = format_kementerian_title(kementerian_name)
    run_t2 = p_title.add_run(f"{formatted_menteri.upper()}\n")
    format_run(run_t2, font_name="Times New Roman", size_pt=14, bold=True)
    
    run_t3 = p_title.add_run("BADAN EKSEKUTIF MAHASISWA\n")
    format_run(run_t3, font_name="Times New Roman", size_pt=14, bold=True)
    
    run_t4 = p_title.add_run("INSTITUT BUDI UTOMO NASIONAL")
    format_run(run_t4, font_name="Times New Roman", size_pt=14, bold=True)
    
    # Logo
    logo_path = get_bem_logo_path()
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(64)
    p_logo.paragraph_format.space_after = Pt(64)
    
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(6.5))
    else:
        run_fallback = p_logo.add_run("[LOGO BEM]")
        format_run(run_fallback, font_name="Times New Roman", size_pt=14, bold=True)
        
    # Bottom Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.line_spacing = 1.15
    p_sub.paragraph_format.space_before = Pt(6)
    p_sub.paragraph_format.space_after = Pt(24)
    
    run_s1 = p_sub.add_run("BADAN EKSEKUTIF MAHASISWA\n")
    format_run(run_s1, font_name="Times New Roman", size_pt=14, bold=True)
    
    run_s2 = p_sub.add_run("INSTITUT BUDI UTOMO NASIONAL MAJALENGKA\n")
    format_run(run_s2, font_name="Times New Roman", size_pt=14, bold=True)
    
    years_str = extract_periode_years(periode_years)
    run_s3 = p_sub.add_run(f"PERIODE {years_str}")
    format_run(run_s3, font_name="Times New Roman", size_pt=14, bold=True)
    
    doc.add_page_break()

def format_run(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        new_borders.append(border)
    tblPr.append(new_borders)

def set_table_widths(table, col_widths):
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = w
                tcPr = row.cells[i]._tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(w.emu / 635)))
                tcW.set(qn('w:type'), 'dxa')
                existing_tcW = tcPr.first_child_found_in("w:tcW")
                if existing_tcW is not None:
                    tcPr.remove(existing_tcW)
                tcPr.append(tcW)

def set_table_indent(table, indent_cm):
    if indent_cm > 0:
        tblPr = table._tbl.tblPr
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).emu / 635)))
        tblInd.set(qn('w:type'), 'dxa')
        existing_ind = tblPr.first_child_found_in('w:tblInd')
        if existing_ind is not None:
            tblPr.remove(existing_ind)
        tblPr.append(tblInd)

def render_table_rows(table, fields, indent_cm=0, bold_label=True, prefix_alpha=False):
    remove_table_borders(table)
    if indent_cm > 0:
        available = 14.0 - indent_cm
        col1 = 4.8
        col2 = 0.5
        col3 = available - col1 - col2
        set_table_widths(table, [Cm(col1), Cm(col2), Cm(col3)])
        set_table_indent(table, indent_cm)
    else:
        set_table_widths(table, [Cm(4.9), Cm(0.7), Cm(8.4)])
    
    for idx, (f_name, f_val) in enumerate(fields):
        row = table.rows[idx]
        
        f_name_text = f_name
        if prefix_alpha:
            f_name_text = f"{chr(ord('a') + idx)}.  {f_name}"
            
        row.cells[0].text = f_name_text
        row.cells[1].text = ":"
        
        cell = row.cells[2]
        if isinstance(f_val, list):
            # Filter out empty items
            clean_items = [clean_proker_name(str(it)) for it in f_val if str(it).strip()]
            clean_items = [it for it in clean_items if it]  # remove blanks after cleaning
            
            p0 = cell.paragraphs[0]
            p0.text = ""
            
            is_single = len(clean_items) == 1
            
            for item_idx, clean_item in enumerate(clean_items):
                if item_idx == 0:
                    p = p0
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                
                if is_single:
                    # Single item — no numbering needed
                    run = p.add_run(clean_item)
                else:
                    # Multiple items — numbered with hanging indent
                    run = p.add_run(f"{item_idx + 1}. {clean_item}")
                    p.paragraph_format.left_indent = Cm(0.6)
                    p.paragraph_format.first_line_indent = Cm(-0.6)
        else:
            cell.text = str(f_val) if f_val else "—"
            
        # Formatting
        for c_idx, cell_item in enumerate(row.cells):
            set_cell_margins(cell_item, top=60, bottom=60, left=100, right=100)
            for p in cell_item.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                if len(cell_item.paragraphs) > 1:
                    p.paragraph_format.space_after = Pt(2)
                else:
                    p.paragraph_format.space_after = Pt(0)
                
                if c_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif c_idx == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for r in p.runs:
                    format_run(r, size_pt=12)
        
        # Hanging indent for label if prefixed with alpha
        p_lbl = row.cells[0].paragraphs[0]
        if prefix_alpha:
            p_lbl.paragraph_format.left_indent = Cm(0.6)
            p_lbl.paragraph_format.first_line_indent = Cm(-0.6)
            
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = bold_label
        for r in row.cells[1].paragraphs[0].runs:
            r.bold = bold_label

def set_document_formatting(doc):
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(4)
        section.bottom_margin = Cm(3)
        section.right_margin = Cm(3)
        
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

def add_document_footer(doc, triwulan_str):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        
        # Main footer
        footer = section.footer
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
            p.text = ""
            
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if triwulan_str.upper() == "MUBESMA":
            txt = "Laporan Pertanggungjawaban Mubesma | Halaman "
        else:
            txt = f"Laporan Pertanggungjawaban Triwulan {triwulan_str} | Halaman "
            
        run_txt = p.add_run(txt)
        format_run(run_txt, font_name="Times New Roman", size_pt=10, italic=True)
        
        run_num = p.add_run()
        format_run(run_num, font_name="Times New Roman", size_pt=10, bold=True)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run_num._r.append(fldChar1)
        run_num._r.append(instrText)
        run_num._r.append(fldChar2)
        
        # Cover page footer (empty)
        first_page_footer = section.first_page_footer
        if not first_page_footer.paragraphs:
            p_first = first_page_footer.add_paragraph()
        else:
            p_first = first_page_footer.paragraphs[0]
            p_first.text = ""

def parse_docx(doc_path):
    if not os.path.exists(doc_path):
        return None
    
    doc = docx.Document(doc_path)
    data = {
        "cover": {"triwulan": "", "kementerian": "", "periode": ""},
        "keanggotaan": {"ketua": "", "sekretaris": "", "bendahara": "", "anggota": ""},
        "keadaan_objektif": "",
        "tugas_pokok": [],
        "fungsi": [],
        "visi": "",
        "misi": [],
        "proker_terlaksana": [],
        "proker_belum_terlaksana": [],
        "anggaran": [],
        "dokumentasi": [],
        "anggaran_summary": {"debet": 0, "kredit": 0, "saldo": 0}
    }
    
    # Parse cover
    cover_text = ""
    for p in doc.paragraphs[:10]:
        cover_text += p.text + "\n"
    
    triwulan_match = re.search(r"(?:TRIWULAN\s+([IVXLCDM\d]+|MUBESMA)|(MUBESMA))", cover_text, re.IGNORECASE)
    kementerian_match = re.search(r"MENTERI\s+([^\n]+)|MENTRI\s+([^\n]+)", cover_text, re.IGNORECASE)
    periode_match = re.search(r"(\d{4}\s*[-–]\s*\d{4})", cover_text)
    
    if triwulan_match:
        val = triwulan_match.group(1) or triwulan_match.group(2) or "MUBESMA"
        data["cover"]["triwulan"] = val.strip().upper()
    elif "laporan pertanggung" in cover_text.lower():
        data["cover"]["triwulan"] = "MUBESMA"
    if kementerian_match:
        kementerian = kementerian_match.group(1) or kementerian_match.group(2)
        data["cover"]["kementerian"] = kementerian.strip()
    if periode_match: data["cover"]["periode"] = periode_match.group(1).strip()
    
    # Parse keadaan objektif
    objektif_header_found = False
    objektif_paras = []
    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        if "keadaan objektif menteri" in p_text.lower() or "keadaan objektif" in p_text.lower() or "pendahuluan" in p_text.lower():
            objektif_header_found = True
            continue
        if objektif_header_found:
            if "keanggotaan" in p_text.lower() or "susunan keanggotaan" in p_text.lower() or p_text.startswith("B. KEANGGOTAAN") or p_text.startswith("II. KEANGGOTAAN") or p_text.startswith("C. REALISASI") or p_text.startswith("III. REALISASI") or "tugas pokok dan fungsi" in p_text.lower() or p_text.startswith("III. TUGAS"):
                break
            if p_text:
                objektif_paras.append(p_text)
    data["keadaan_objektif"] = "\n\n".join(objektif_paras)

    # Parse Tugas Pokok dan Fungsi (Only for MUBESMA)
    tupoksi_header_found = False
    current_list_type = None
    tugas_list = []
    fungsi_list = []
    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        if "tugas pokok dan fungsi" in p_text.lower() or p_text.startswith("III. TUGAS"):
            tupoksi_header_found = True
            continue
        if tupoksi_header_found:
            if "evaluasi pencapaian" in p_text.lower() or p_text.startswith("IV. EVALUASI") or "realisasi program kerja" in p_text.lower() or p_text.startswith("V. REALISASI") or p_text.startswith("C. REALISASI"):
                break
            if "tugas pokok" in p_text.lower() and (p_text.startswith("A.") or p_text.startswith("A. ")):
                current_list_type = 'tugas'
                continue
            elif "fungsi" in p_text.lower() and (p_text.startswith("B.") or p_text.startswith("B. ")):
                current_list_type = 'fungsi'
                continue
            if current_list_type and p_text:
                cleaned_item = re.sub(r'^\d+[\.\s-]*', '', p_text).strip()
                if cleaned_item:
                    if current_list_type == 'tugas':
                        tugas_list.append(cleaned_item)
                    elif current_list_type == 'fungsi':
                        fungsi_list.append(cleaned_item)
    data["tugas_pokok"] = tugas_list
    data["fungsi"] = fungsi_list

    # Parse Visi and Misi (Only for MUBESMA)
    visi_hdr_found = False
    current_list_type = None
    visi_txt = ""
    misi_list = []
    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        if "evaluasi pencapaian" in p_text.lower() or p_text.startswith("IV. EVALUASI"):
            visi_hdr_found = True
            continue
        if visi_hdr_found:
            if "realisasi program kerja" in p_text.lower() or p_text.startswith("V. REALISASI") or p_text.startswith("C. REALISASI"):
                break
            if "visi" in p_text.lower() and (p_text.startswith("A.") or p_text.startswith("A. ")):
                current_list_type = 'visi'
                continue
            elif "misi" in p_text.lower() and (p_text.startswith("B.") or p_text.startswith("B. ")):
                current_list_type = 'misi'
                continue
            if current_list_type == 'visi' and p_text and not p_text.startswith("Visi dan Misi"):
                visi_txt = p_text
            elif current_list_type == 'misi' and p_text:
                cleaned_item = re.sub(r'^\d+[\.\s-]*', '', p_text).strip()
                if cleaned_item:
                    misi_list.append(cleaned_item)
    data["visi"] = visi_txt
    data["misi"] = misi_list

    # Walk through tables sequentially
    current_proker = None
    for table in doc.tables:
        row0_text = [cell.text.strip().lower() for cell in table.rows[0].cells]
        first_col_texts = [row.cells[0].text.strip().lower() for row in table.rows]
        
        # 1. Check for Keanggotaan table
        if any("ketua menteri" in t for t in first_col_texts):
            for row in table.rows:
                if len(row.cells) >= 3:
                    f_name = row.cells[0].text.strip().lower()
                    f_val = row.cells[2].text.strip()
                    if "ketua" in f_name:
                        data["keanggotaan"]["ketua"] = f_val
                    elif "sekretaris" in f_name:
                        data["keanggotaan"]["sekretaris"] = f_val
                    elif "bendahara" in f_name:
                        data["keanggotaan"]["bendahara"] = f_val
                    elif "anggota" in f_name:
                        data["keanggotaan"]["anggota"] = f_val
            continue

        # 2. Check for Budget table (has Tanggal, Keterangan, Saldo)
        if "tanggal" in row0_text and "keterangan" in row0_text and "saldo" in row0_text:
            if "kementerian" in row0_text:
                row_last = table.rows[-1]
                row_last_cells = [cell.text.strip() for cell in row_last.cells]
                if "grand total" in row_last_cells[0].lower() or "total" in row_last_cells[0].lower():
                    data["anggaran_summary"]["debet"] = clean_currency(row_last_cells[1])
                    data["anggaran_summary"]["kredit"] = clean_currency(row_last_cells[2])
                    data["anggaran_summary"]["saldo"] = clean_currency(row_last_cells[3])
                continue
                
            col_indices = {}
            for h in ["tanggal", "keterangan", "uraian", "debet", "kredit", "saldo"]:
                for idx, hl in enumerate(row0_text):
                    if h in hl:
                        col_indices[h] = idx
                        break
            
            proker_anggaran = []
            proker_summary = {"debet": 0, "kredit": 0, "saldo": 0}
            for r_idx in range(1, len(table.rows)):
                row_cells = [cell.text.strip() for cell in table.rows[r_idx].cells]
                if "total" in row_cells[0].lower() or "total" in row_cells[1].lower():
                    proker_summary["debet"] = clean_currency(row_cells[col_indices["debet"]])
                    proker_summary["kredit"] = clean_currency(row_cells[col_indices["kredit"]])
                    proker_summary["saldo"] = clean_currency(row_cells[col_indices["saldo"]])
                    continue
                
                tx = {
                    "tanggal": row_cells[col_indices["tanggal"]],
                    "keterangan": row_cells[col_indices["keterangan"]],
                    "uraian": row_cells[col_indices["uraian"]] if "uraian" in col_indices else "",
                    "debet": row_cells[col_indices["debet"]],
                    "kredit": row_cells[col_indices["kredit"]],
                    "saldo": row_cells[col_indices["saldo"]]
                }
                proker_anggaran.append(tx)
                data["anggaran"].append(tx)
            
            if current_proker:
                current_proker["tidak_menggunakan_anggaran"] = False
                current_proker["anggaran"] = proker_anggaran
                current_proker["anggaran_summary"] = proker_summary
            continue

        # 3. Check for Proker table (Realized or Unrealized)
        if any("nama program kerja" in t or "nama kegiatan" in t for t in first_col_texts):
            fields = {}
            for row in table.rows:
                if len(row.cells) >= 3:
                    f_name = row.cells[0].text.strip()
                    # Clean prefix like "a.  ", "b.  ", "1. ", etc.
                    clean_key = re.sub(r'^[a-zA-Z0-9]+[\.\s]+', '', f_name).strip()
                    f_val = row.cells[2].text.strip()
                    fields[clean_key] = f_val
            
            has_evaluasi = any("evaluasi" in k.lower() for k in fields.keys())
            if has_evaluasi:
                current_proker = {
                    "Nama Program Kerja": fields.get("Nama Program Kerja", fields.get("Nama Kegiatan", "")),
                    "Tempat Kegiatan": fields.get("Tempat Kegiatan", fields.get("Tempat", "")),
                    "Sifat": fields.get("Sifat", "Internal"),
                    "Tema Kegiatan": fields.get("Tema Kegiatan", ""),
                    "Tujuan": fields.get("Tujuan", fields.get("Tujuan Kegiatan", "")),
                    "Tanggal Kegiatan": fields.get("Tanggal Kegiatan", ""),
                    "Penanggung Jawab": fields.get("Penanggung Jawab", ""),
                    "Peserta Kegiatan": fields.get("Peserta Kegiatan", ""),
                    "Evaluasi": fields.get("Evaluasi & Saran", fields.get("Evaluasi", "")),
                    "tidak_menggunakan_anggaran": True,
                    "anggaran": [],
                    "anggaran_summary": {"debet": 0, "kredit": 0, "saldo": 0},
                    "dokumentasi": [],
                    "nota_belanja": []
                }
                data["proker_terlaksana"].append(current_proker)
            else:
                data["proker_belum_terlaksana"].append(fields)
            continue
            
        # 4. Check for Documentation / Nota Belanja table (grid of photos: 2 columns, borderless)
        if len(table.rows) > 1 and len(table.columns) == 2:
            preceding_text = ""
            try:
                parent_elm = doc.element.body
                body_children = list(parent_elm.iterchildren())
                tbl_idx = body_children.index(table._element)
                for idx in range(tbl_idx - 1, -1, -1):
                    child = body_children[idx]
                    if child.tag.endswith('p'): # It is a paragraph
                        p = docx.text.paragraph.Paragraph(child, doc)
                        p_text = p.text.strip()
                        if p_text:
                            preceding_text = p_text
                            break
            except Exception:
                pass
            
            is_nota = "nota" in preceding_text.lower()
            
            proker_docs = []
            for r_idx in range(0, len(table.rows), 2):
                if r_idx + 1 < len(table.rows):
                    cap_row = table.rows[r_idx + 1]
                    for c_idx in range(2):
                        if c_idx < len(cap_row.cells):
                            cap_text = cap_row.cells[c_idx].text.strip()
                            if cap_text:
                                xml_str = table.rows[r_idx].cells[c_idx]._tc.xml
                                has_image = "w:drawing" in xml_str or "w:graphic" in xml_str
                                doc_item = {
                                    "caption": cap_text,
                                    "has_image": has_image
                                }
                                proker_docs.append(doc_item)
                                if not is_nota:
                                    data["dokumentasi"].append(doc_item)
            if current_proker:
                if is_nota:
                    current_proker["nota_belanja"] = proker_docs
                else:
                    current_proker["dokumentasi"] = proker_docs
            continue
            
    # Parse evaluasi kinerja pribadi
    eval_pribadi_header_found = False
    eval_pribadi_paras = []
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not eval_pribadi_header_found and ("evaluasi kinerja pribadi" in p_text.lower() or "evaluasi kinerja menteri" in p_text.lower()):
            eval_pribadi_header_found = True
            continue
        if eval_pribadi_header_found:
            # stop when we reach Section IX / Penutup or Section F
            if "penutup" in p_text.lower() or "evaluasi anggota dan internal" in p_text.lower() or p_text.startswith("IX.") or p_text.startswith("X.") or p_text.startswith("F.") or p_text.startswith("G.") or "ringkasan anggaran" in p_text.lower():
                break
            if p_text:
                eval_pribadi_paras.append(p_text)
    data["evaluasi_kinerja_pribadi"] = "\n\n".join(eval_pribadi_paras)

    # Parse evaluasi anggota
    eval_anggota_list = []
    for table in doc.tables:
        row0_text = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "nama anggota" in row0_text and "kepribadian" in row0_text and "kinerja" in row0_text:
            col_indices = {}
            for h in ["nama anggota", "kepribadian", "kinerja"]:
                for idx, hl in enumerate(row0_text):
                    if h in hl:
                        col_indices[h] = idx
                        break
            for r_idx in range(1, len(table.rows)):
                row_cells = [cell.text.strip() for cell in table.rows[r_idx].cells]
                if len(row_cells) > max(col_indices.values()):
                    eval_anggota_list.append({
                        "nama": row_cells[col_indices["nama anggota"]],
                        "kepribadian": row_cells[col_indices["kepribadian"]],
                        "kinerja": row_cells[col_indices["kinerja"]]
                    })
    if eval_anggota_list:
        data["evaluasi_anggota_internal"] = eval_anggota_list

    # For global backward compatibility, let's aggregate anggaran summary from all prokers if not already set by BEM summary table
    if data["anggaran"] and data["anggaran_summary"]["debet"] == 0:
        total_deb = 0
        total_kred = 0
        running = 0
        for tx in data["anggaran"]:
            deb = clean_currency(tx["debet"])
            kred = clean_currency(tx["kredit"])
            total_deb += deb
            total_kred += kred
            running = running + deb - kred
        data["anggaran_summary"]["debet"] = total_deb
        data["anggaran_summary"]["kredit"] = total_kred
        data["anggaran_summary"]["saldo"] = running
        
    return data

def run_validation(doc_path):
    res = parse_docx(doc_path)
    if not res:
        return {"status": "ERROR", "message": f"File tidak ditemukan: {doc_path}"}
    
    checklist = {
        "cover": bool(res["cover"]["triwulan"] and res["cover"]["kementerian"] and res["cover"]["periode"]),
        "keanggotaan": bool(res["keanggotaan"]["ketua"] and res["keanggotaan"]["sekretaris"] and res["keanggotaan"]["bendahara"]),
        "proker_terlaksana": len(res["proker_terlaksana"]) > 0,
        "tabel_anggaran_kolom": True,
        "tabel_anggaran_saldo": True,
        "tabel_anggaran_total": True,
        "proker_belum_terlaksana": len(res["proker_belum_terlaksana"]) > 0,
        "dokumentasi": True
    }
    
    errors = []
    
    if not checklist["cover"]:
        errors.append("Cover tidak lengkap (pastikan mencakup Triwulan, Kementerian, dan Periode BEM).")
    if not checklist["keanggotaan"]:
        errors.append("Keanggotaan tidak lengkap (Ketua, Sekretaris, Bendahara wajib diisi).")
    if not checklist["proker_terlaksana"]:
        errors.append("Minimal harus ada 1 Program Kerja Terlaksana.")
    else:
        # Check proker fields
        for pk in res["proker_terlaksana"]:
            name = pk.get("Nama Program Kerja", pk.get("Nama Kegiatan", "Tanpa Nama"))
            required = ["Nama Program Kerja", "Sifat", "Tema Kegiatan", "Tujuan", "Tanggal Kegiatan", "Penanggung Jawab", "Peserta Kegiatan", "Evaluasi"]
            missing = []
            for r in required:
                found = False
                for k in pk.keys():
                    if r.lower() in k.lower():
                        found = True
                        if not pk[k] or pk[k] == "-":
                            missing.append(r)
                        break
                if not found:
                    missing.append(r)
            if missing:
                errors.append(f"Proker Terlaksana '{name}' tidak lengkap, field kosong: {', '.join(missing)}")
                checklist["proker_terlaksana"] = False
                
            # Validate proker budget if it has one
            if not pk.get("tidak_menggunakan_anggaran", False) and pk.get("anggaran"):
                running = 0
                has_error = False
                for idx, row in enumerate(pk["anggaran"]):
                    deb = clean_currency(row["debet"])
                    kred = clean_currency(row["kredit"])
                    sal_doc = clean_currency(row["saldo"])
                    
                    if idx == 0:
                        running = deb - kred
                    else:
                        running = running + deb - kred
                    
                    if running != sal_doc:
                        has_error = True
                        errors.append(f"Baris anggaran proker '{name}' baris {idx+1} ({row['tanggal']} - {row['keterangan']}): Saldo dokumen = {format_currency(sal_doc)}, seharusnya {format_currency(running)}")
                if has_error:
                    checklist["tabel_anggaran_saldo"] = False
                    
                sum_deb = sum(clean_currency(r["debet"]) for r in pk["anggaran"])
                sum_kred = sum(clean_currency(r["kredit"]) for r in pk["anggaran"])
                p_summary = pk.get("anggaran_summary", {"debet": 0, "kredit": 0, "saldo": 0})
                
                if p_summary["debet"] != sum_deb or p_summary["kredit"] != sum_kred or p_summary["saldo"] != running:
                    checklist["tabel_anggaran_total"] = False
                    errors.append(f"Baris TOTAL anggaran proker '{name}' salah! Dokumen: (Debet={format_currency(p_summary['debet'])}, Kredit={format_currency(p_summary['kredit'])}, Saldo={format_currency(p_summary['saldo'])}). Seharusnya: (Debet={format_currency(sum_deb)}, Kredit={format_currency(sum_kred)}, Saldo={format_currency(running)})")
            
            # Check documentation captions
            if pk.get("dokumentasi"):
                for p_idx, doc_item in enumerate(pk["dokumentasi"]):
                    if not doc_item.get("caption"):
                        errors.append(f"Foto dokumentasi ke-{p_idx+1} pada proker '{name}' tidak memiliki caption.")
                        checklist["dokumentasi"] = False
                        
            # Check nota belanja captions
            if pk.get("nota_belanja"):
                for p_idx, doc_item in enumerate(pk["nota_belanja"]):
                    if not doc_item.get("caption"):
                        errors.append(f"Foto nota belanja ke-{p_idx+1} pada proker '{name}' tidak memiliki caption.")
                        
    status = "LENGKAP" if all(checklist.values()) else "TIDAK LENGKAP"
    
    return {
        "status": status,
        "checklist": checklist,
        "errors": errors,
        "details": res
    }

def generate_lpj(output_path, config_data):
    doc = docx.Document()
    set_document_formatting(doc)
    
    triwulan_str = config_data['cover']['triwulan'].upper()
    kementrian_str = config_data['cover']['kementerian']
    years_str = config_data['cover']['periode']
    
    is_mubesma = (triwulan_str == "MUBESMA")
    pref_a = "I." if is_mubesma else "A."
    pref_b = "II." if is_mubesma else "B."
    pref_c = "VI." if is_mubesma else "C."
    pref_d = "VII." if is_mubesma else "D."
    pref_e = "VIII." if is_mubesma else "E."
    pref_f = "IX." if is_mubesma else "F."
    pref_penutup = "IX." if is_mubesma else "D."
    
    if triwulan_str == "MUBESMA":
        add_minister_cover(doc, kementrian_str, years_str)
    else:
        # Header Title Block (No page break, matches Image 3 exactly)
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.line_spacing = 1.0
        cover_title.paragraph_format.space_before = Pt(0)
        cover_title.paragraph_format.space_after = Pt(12)
        
        kementrian_formatted = format_kementerian_title(kementrian_str)
        years_extracted = extract_periode_years(years_str)
        
        run1 = cover_title.add_run(f"LAPORAN PERTANGGUNG JAWABAN TRIWULAN {triwulan_str}\n")
        format_run(run1, size_pt=12, bold=True)
        
        run2 = cover_title.add_run(f"{kementrian_formatted}\n")
        format_run(run2, size_pt=12, bold=True)
        
        run3 = cover_title.add_run(f"BEM INSTBUNAS MAJALENGKA {years_extracted}")
        format_run(run3, size_pt=12, bold=True)
    
    # A. KEADAAN OBJEKTIF MENTERI / PENDAHULUAN
    p_hdr_a = doc.add_paragraph()
    p_hdr_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr_a.paragraph_format.space_before = Pt(12)
    p_hdr_a.paragraph_format.space_after = Pt(6)
    p_hdr_a.paragraph_format.keep_with_next = True
    title_a = "I. PENDAHULUAN" if is_mubesma else "A. KEADAAN OBJEKTIF MENTERI"
    format_run(p_hdr_a.add_run(title_a), size_pt=12, bold=True)
    
    obj_text = config_data.get("keadaan_objektif", "").strip()
    if not obj_text:
        obj_text = "Deskripsi Keadaan Objektif..."
    
    for line in obj_text.split("\n"):
        line_clean = line.strip()
        if not line_clean:
            doc.add_paragraph()
            continue
        p_fungsi = doc.add_paragraph()
        p_fungsi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fungsi.paragraph_format.line_spacing = 1.5
        p_fungsi.paragraph_format.space_before = Pt(0)
        p_fungsi.paragraph_format.space_after = Pt(0)
        p_fungsi.paragraph_format.left_indent = Cm(0.5)
        p_fungsi.paragraph_format.first_line_indent = Cm(1.0)
        format_run(p_fungsi.add_run(line_clean), size_pt=12)
    
    # B. SUSUNAN KEANGGOTAAN (flow on same page — no page break)
    p_hdr_b = doc.add_paragraph()
    p_hdr_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr_b.paragraph_format.space_before = Pt(12)
    p_hdr_b.paragraph_format.space_after = Pt(6)
    p_hdr_b.paragraph_format.keep_with_next = True
    format_run(p_hdr_b.add_run(f"{pref_b} SUSUNAN KEANGGOTAAN"), size_pt=12, bold=True)
    
    anggota_val = config_data['keanggotaan'].get('anggota', '')
    has_anggota = False
    if anggota_val:
        if isinstance(anggota_val, list):
            anggota_str = ", ".join(anggota_val).strip()
        else:
            anggota_str = str(anggota_val).strip()
        if anggota_str and anggota_str != "-" and anggota_str.lower() != "null":
            has_anggota = True
            
    rows_data = [
        ("Ketua Menteri", config_data['keanggotaan'].get('ketua', '')),
        ("Sekretaris", config_data['keanggotaan'].get('sekretaris', '')),
        ("Bendahara", config_data['keanggotaan'].get('bendahara', ''))
    ]
    if has_anggota:
        rows_data.append(("Anggota", anggota_str))
        
    table_k = doc.add_table(rows=len(rows_data), cols=3)
    render_table_rows(table_k, rows_data, indent_cm=0.5, bold_label=False, prefix_alpha=False)
    
    if is_mubesma:
        # III. TUGAS POKOK DAN FUNGSI
        p_hdr_tup = doc.add_paragraph()
        p_hdr_tup.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_tup.paragraph_format.space_before = Pt(12)
        p_hdr_tup.paragraph_format.space_after = Pt(6)
        p_hdr_tup.paragraph_format.keep_with_next = True
        format_run(p_hdr_tup.add_run("III. TUGAS POKOK DAN FUNGSI"), size_pt=12, bold=True)
        
        # A. Tugas Pokok
        p_sub_tp = doc.add_paragraph()
        p_sub_tp.paragraph_format.space_before = Pt(6)
        p_sub_tp.paragraph_format.space_after = Pt(4)
        p_sub_tp.paragraph_format.keep_with_next = True
        p_sub_tp.paragraph_format.left_indent = Cm(0.5)
        format_run(p_sub_tp.add_run("A. Tugas Pokok"), size_pt=12, bold=True)
        
        tugas_list = config_data.get("tugas_pokok", [])
        if not tugas_list:
            tugas_list = ["(Belum diatur)"]
        for t_idx, t_item in enumerate(tugas_list, 1):
            p_t_item = doc.add_paragraph()
            p_t_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_t_item.paragraph_format.line_spacing = 1.15
            p_t_item.paragraph_format.space_after = Pt(4)
            p_t_item.paragraph_format.left_indent = Cm(1.5)
            p_t_item.paragraph_format.first_line_indent = Cm(-0.5)
            format_run(p_t_item.add_run(f"{t_idx}.\t{t_item}"), size_pt=12)
            
        # B. Fungsi
        p_sub_f = doc.add_paragraph()
        p_sub_f.paragraph_format.space_before = Pt(6)
        p_sub_f.paragraph_format.space_after = Pt(4)
        p_sub_f.paragraph_format.keep_with_next = True
        p_sub_f.paragraph_format.left_indent = Cm(0.5)
        format_run(p_sub_f.add_run("B. Fungsi"), size_pt=12, bold=True)
        
        fungsi_list = config_data.get("fungsi", [])
        if not fungsi_list:
            fungsi_list = ["(Belum diatur)"]
        for f_idx, f_item in enumerate(fungsi_list, 1):
            p_f_item = doc.add_paragraph()
            p_f_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_f_item.paragraph_format.line_spacing = 1.15
            p_f_item.paragraph_format.space_after = Pt(4)
            p_f_item.paragraph_format.left_indent = Cm(1.5)
            p_f_item.paragraph_format.first_line_indent = Cm(-0.5)
            format_run(p_f_item.add_run(f"{f_idx}.\t{f_item}"), size_pt=12)
            
        # IV. EVALUASI PENCAPAIAN VISI DAN MISI
        p_hdr_vm = doc.add_paragraph()
        p_hdr_vm.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_vm.paragraph_format.space_before = Pt(12)
        p_hdr_vm.paragraph_format.space_after = Pt(6)
        p_hdr_vm.paragraph_format.keep_with_next = True
        format_run(p_hdr_vm.add_run("IV. EVALUASI PENCAPAIAN VISI DAN MISI"), size_pt=12, bold=True)
        
        p_sub_vm = doc.add_paragraph()
        p_sub_vm.paragraph_format.space_before = Pt(6)
        p_sub_vm.paragraph_format.space_after = Pt(6)
        p_sub_vm.paragraph_format.keep_with_next = True
        p_sub_vm.paragraph_format.left_indent = Cm(0.5)
        org_name = "Badan Perwakilan Mahasiswa" if "bpm" in output_path.lower() else "Badan Eksekutif Mahasiswa"
        format_run(p_sub_vm.add_run(f"Visi dan Misi {org_name} INSTBUNAS Majalengka"), size_pt=12, bold=True)
        
        # A. Visi
        p_v_hdr = doc.add_paragraph()
        p_v_hdr.paragraph_format.space_before = Pt(6)
        p_v_hdr.paragraph_format.space_after = Pt(4)
        p_v_hdr.paragraph_format.keep_with_next = True
        p_v_hdr.paragraph_format.left_indent = Cm(0.5)
        format_run(p_v_hdr.add_run("A. Visi"), size_pt=12, bold=True)
        
        visi_text = config_data.get("visi", "").strip()
        if not visi_text:
            visi_text = "(Belum diatur)"
        p_v_txt = doc.add_paragraph()
        p_v_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_v_txt.paragraph_format.line_spacing = 1.15
        p_v_txt.paragraph_format.space_after = Pt(12)
        p_v_txt.paragraph_format.left_indent = Cm(0.5)
        format_run(p_v_txt.add_run(visi_text), size_pt=12)
        
        # B. Misi
        p_m_hdr = doc.add_paragraph()
        p_m_hdr.paragraph_format.space_before = Pt(6)
        p_m_hdr.paragraph_format.space_after = Pt(4)
        p_m_hdr.paragraph_format.keep_with_next = True
        p_m_hdr.paragraph_format.left_indent = Cm(0.5)
        format_run(p_m_hdr.add_run("B. Misi"), size_pt=12, bold=True)
        
        misi_list = config_data.get("misi", [])
        if not misi_list:
            misi_list = ["(Belum diatur)"]
        for m_idx, m_item in enumerate(misi_list, 1):
            p_m_item = doc.add_paragraph()
            p_m_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_m_item.paragraph_format.line_spacing = 1.15
            p_m_item.paragraph_format.space_after = Pt(4)
            p_m_item.paragraph_format.left_indent = Cm(1.5)
            p_m_item.paragraph_format.first_line_indent = Cm(-0.5)
            format_run(p_m_item.add_run(f"{m_idx}.\t{m_item}"), size_pt=12)
            
    # Sort proker terlaksana chronologically
    proker_terlaksana = config_data.get("proker_terlaksana", [])
    proker_terlaksana.sort(key=lambda x: parse_id_date(x.get('Tanggal Kegiatan', '')))

    if is_mubesma:
        # V. PROGRAM KERJA (Summary Table)
        p_hdr_v = doc.add_paragraph()
        p_hdr_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_v.paragraph_format.space_before = Pt(12)
        p_hdr_v.paragraph_format.space_after = Pt(6)
        p_hdr_v.paragraph_format.keep_with_next = True
        format_run(p_hdr_v.add_run("V. PROGRAM KERJA"), size_pt=12, bold=True)
        
        table_v = doc.add_table(rows=1, cols=5)
        table_v.style = 'Table Grid'
        set_table_indent(table_v, 0.5)
        set_table_widths(table_v, [Cm(1.0), Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.0)])
        
        headers_v = ["NO", "WAKTU", "PROGRAM KERJA", "NAMA KEGIATAN", "TEMPAT"]
        hdr_cells = table_v.rows[0].cells
        for c_idx, h in enumerate(headers_v):
            hdr_cells[c_idx].text = h
            set_cell_shading(hdr_cells[c_idx], "D3D3D3")
            set_cell_margins(hdr_cells[c_idx], top=100, bottom=100, left=120, right=120)
            hdr_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[c_idx].paragraphs[0].paragraph_format.line_spacing = 1.0
            format_run(hdr_cells[c_idx].paragraphs[0].runs[0], size_pt=12, bold=True)
            
        proker_groups_v = []
        prev_prog = None
        current_no = 1
        for pk_row in proker_terlaksana:
            prog_name = clean_proker_name(pk_row.get('Nama Program Kerja', '—'))
            if prev_prog == prog_name and proker_groups_v:
                proker_groups_v[-1]['rows'].append(pk_row)
            else:
                proker_groups_v.append({
                    'name': pk_row.get('Nama Program Kerja', '—'),
                    'start_no': current_no,
                    'rows': [pk_row]
                })
                prev_prog = prog_name
            current_no += 1
            
        row_idx = 1
        for grp in proker_groups_v:
            span = len(grp['rows'])
            start_row = row_idx
            for r_idx, r_row in enumerate(grp['rows']):
                row_cells = table_v.add_row().cells
                set_table_widths(table_v, [Cm(1.0), Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.0)])
                
                if r_idx == 0:
                    row_cells[0].text = str(grp['start_no'])
                    row_cells[2].text = grp['name']
                
                row_cells[1].text = format_id_date_str(r_row.get('Tanggal Kegiatan', '—'))
                row_cells[3].text = r_row.get('Nama Kegiatan', r_row.get('Nama Program Kerja', '—'))
                row_cells[4].text = r_row.get('Tempat Kegiatan', r_row.get('Tempat', '—'))
                
                for c_idx, cell in enumerate(row_cells):
                    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    cell.vertical_alignment = 0
                    p = cell.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if c_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for r in p.runs:
                        format_run(r, size_pt=12)
                row_idx += 1
                
            end_row = row_idx - 1
            if span > 1:
                cell_no_start = table_v.cell(start_row, 0)
                cell_no_end = table_v.cell(end_row, 0)
                cell_no_start.merge(cell_no_end)
                cell_no_start.vertical_alignment = 0
                
                cell_pk_start = table_v.cell(start_row, 2)
                cell_pk_end = table_v.cell(end_row, 2)
                cell_pk_start.merge(cell_pk_end)
                cell_pk_start.vertical_alignment = 0

    # C. PROGRAM KERJA YANG TEREALISASI
    p_hdr_c = doc.add_paragraph()
    p_hdr_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr_c.paragraph_format.space_before = Pt(12)
    p_hdr_c.paragraph_format.space_after = Pt(6)
    p_hdr_c.paragraph_format.keep_with_next = True
    format_run(p_hdr_c.add_run(f"{pref_c} PROGRAM KERJA YANG TEREALISASI"), size_pt=12, bold=True)
    
    INDENT_PROKER = 1.0  # cm indent for proker sub-content
    
    for idx, pk in enumerate(proker_terlaksana, 1):
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_before = Pt(12)
        p_name.paragraph_format.space_after = Pt(6)
        p_name.paragraph_format.keep_with_next = True
        p_name.paragraph_format.left_indent = Cm(0.5)
        
        name_str = clean_proker_name(pk.get('Nama Program Kerja', 'Proker'))
        format_run(p_name.add_run(f"{idx}. {name_str}"), size_pt=12, bold=True)
        
        table = doc.add_table(rows=9, cols=3)
        fields = [
            ("Nama Kegiatan", clean_proker_name(pk.get("Nama Kegiatan", pk.get("Nama Program Kerja", "")))),
            ("Tempat Kegiatan", pk.get("Tempat Kegiatan", pk.get("Tempat", ""))),
            ("Sifat", pk.get("Sifat", "Internal")),
            ("Tema Kegiatan", pk.get("Tema Kegiatan", "")),
            ("Tujuan", pk.get("Tujuan", "")),
            ("Tanggal Kegiatan", format_id_date_str(pk.get("Tanggal Kegiatan", ""))),
            ("Penanggung Jawab", pk.get("Penanggung Jawab", "")),
            ("Peserta Kegiatan", pk.get("Peserta Kegiatan", "")),
            ("Evaluasi & Saran", pk.get("Evaluasi & Saran", pk.get("Evaluasi", "")))
        ]
        render_table_rows(table, fields, indent_cm=INDENT_PROKER, bold_label=False, prefix_alpha=True)
        
        # Sub-bagian: Realisasi Anggaran
        p_sub_ang = doc.add_paragraph()
        p_sub_ang.paragraph_format.space_before = Pt(12)
        p_sub_ang.paragraph_format.space_after = Pt(6)
        p_sub_ang.paragraph_format.keep_with_next = True
        p_sub_ang.paragraph_format.left_indent = Cm(INDENT_PROKER)
        format_run(p_sub_ang.add_run("j.  Realisasi Anggaran"), size_pt=12, bold=False)
        
        tidak_menggunakan_anggaran = pk.get("tidak_menggunakan_anggaran", False)
        anggaran_list = pk.get("anggaran", [])
        
        if tidak_menggunakan_anggaran or not anggaran_list:
            p_ang_empty = doc.add_paragraph()
            p_ang_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_ang_empty.add_run("(Tidak ada realisasi anggaran)"), size_pt=12, italic=True)
        else:
            table_ang = doc.add_table(rows=1, cols=6)
            set_table_indent(table_ang, INDENT_PROKER)
            table_ang.style = 'Table Grid'
            set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
            
            headers = ["Tanggal", "Keterangan", "Uraian", "Debet", "Kredit", "Saldo"]
            hdr_cells = table_ang.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                set_cell_shading(hdr_cells[i], "D3D3D3")
                set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
                format_run(hdr_cells[i].paragraphs[0].runs[0], size_pt=12, bold=True)
                
            running_balance = 0
            total_debet = 0
            total_kredit = 0
            
            for tx in anggaran_list:
                row_cells = table_ang.add_row().cells
                set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
                
                deb_val = clean_currency(tx.get("debet", 0))
                kred_val = clean_currency(tx.get("kredit", 0))
                
                total_debet += deb_val
                total_kredit += kred_val
                
                if running_balance == 0 and len(table_ang.rows) == 2:
                    running_balance = deb_val - kred_val
                else:
                    running_balance = running_balance + deb_val - kred_val
                    
                row_cells[0].text = tx.get("tanggal", "")
                row_cells[1].text = tx.get("keterangan", "")
                row_cells[2].text = tx.get("uraian", "")
                row_cells[3].text = format_currency(deb_val) if deb_val > 0 else ""
                row_cells[4].text = format_currency(kred_val) if kred_val > 0 else ""
                row_cells[5].text = format_currency(running_balance)
                
                for idx_c, cell in enumerate(row_cells):
                    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    p = cell.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if idx_c in [3, 4, 5]:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        format_run(r, size_pt=12)
                        
            tot_cells = table_ang.add_row().cells
            set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
            
            tot_cells[0].text = "TOTAL"
            tot_cells[1].text = "TOTAL"
            tot_cells[2].text = "TOTAL"
            tot_cells[3].text = format_currency(total_debet)
            tot_cells[4].text = format_currency(total_kredit)
            tot_cells[5].text = format_currency(running_balance)
            
            for idx_c, cell in enumerate(tot_cells):
                set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                set_cell_shading(cell, "EAEAEA")
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if idx_c in [3, 4, 5]:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    format_run(r, size_pt=12, bold=True)
                    
        # Sub-bagian: Dokumentasi Kegiatan
        p_sub_dok = doc.add_paragraph()
        p_sub_dok.paragraph_format.space_before = Pt(12)
        p_sub_dok.paragraph_format.space_after = Pt(6)
        p_sub_dok.paragraph_format.keep_with_next = True
        p_sub_dok.paragraph_format.left_indent = Cm(INDENT_PROKER)
        format_run(p_sub_dok.add_run("k.  Dokumentasi Kegiatan"), size_pt=12, bold=False)
        
        doc_list = pk.get("dokumentasi", [])
        num_photos = len(doc_list)
        
        if num_photos == 0:
            p_dok_empty = doc.add_paragraph()
            p_dok_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_dok_empty.add_run("(Dokumentasi tidak tersedia)"), size_pt=12, italic=True)
        else:
            num_rows_doc = ((num_photos + 1) // 2) * 2
            table_doc = doc.add_table(rows=num_rows_doc, cols=2)
            set_table_indent(table_doc, INDENT_PROKER)
            remove_table_borders(table_doc)
            set_table_widths(table_doc, [Cm(6.5), Cm(6.5)])
            
            for i, photo in enumerate(doc_list):
                grid_row = (i // 2) * 2
                grid_col = i % 2
                
                img_cell = table_doc.rows[grid_row].cells[grid_col]
                set_cell_margins(img_cell, top=80, bottom=40, left=100, right=100)
                p_img = img_cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.line_spacing = 1.0
                p_img.paragraph_format.space_before = Pt(0)
                p_img.paragraph_format.space_after = Pt(0)
                
                photo_path = resolve_photo_path(photo.get("file_path", ""))
                compat_path, is_temp = get_docx_compatible_image(photo_path)
                if compat_path:
                    try:
                        p_img.add_run().add_picture(compat_path, width=Cm(6))
                    except Exception as e:
                        print(f"Exception adding picture: {e}")
                        run_pl = p_img.add_run(f"[Foto: {photo.get('caption', 'Kegiatan')}]")
                        format_run(run_pl, size_pt=11, italic=True)
                    finally:
                        if is_temp and os.path.exists(compat_path):
                            try:
                                os.remove(compat_path)
                            except:
                                pass
                else:
                    run_pl = p_img.add_run("[Foto Kegiatan]")
                    format_run(run_pl, size_pt=11, italic=True)
                    
                cap_cell = table_doc.rows[grid_row + 1].cells[grid_col]
                set_cell_margins(cap_cell, top=40, bottom=80, left=100, right=100)
                p_cap = cap_cell.paragraphs[0]
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.line_spacing = 1.0
                p_cap.paragraph_format.space_before = Pt(0)
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(photo.get("caption", ""))
                format_run(run_cap, size_pt=11, italic=True)
        # Sub-bagian: Nota Belanja
        p_sub_nota = doc.add_paragraph()
        p_sub_nota.paragraph_format.space_before = Pt(12)
        p_sub_nota.paragraph_format.space_after = Pt(6)
        p_sub_nota.paragraph_format.keep_with_next = True
        p_sub_nota.paragraph_format.left_indent = Cm(INDENT_PROKER)
        format_run(p_sub_nota.add_run("l.  Nota Belanja"), size_pt=12, bold=False)
        
        nota_list = pk.get("nota_belanja", [])
        num_notas = len(nota_list)
        
        if num_notas == 0:
            p_nota_empty = doc.add_paragraph()
            p_nota_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_nota_empty.add_run("(Nota belanja tidak tersedia)"), size_pt=12, italic=True)
        else:
            num_rows_nota = ((num_notas + 1) // 2) * 2
            table_nota = doc.add_table(rows=num_rows_nota, cols=2)
            set_table_indent(table_nota, INDENT_PROKER)
            remove_table_borders(table_nota)
            set_table_widths(table_nota, [Cm(6.5), Cm(6.5)])
            
            for i, photo in enumerate(nota_list):
                grid_row = (i // 2) * 2
                grid_col = i % 2
                
                img_cell = table_nota.rows[grid_row].cells[grid_col]
                set_cell_margins(img_cell, top=80, bottom=40, left=100, right=100)
                p_img = img_cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.line_spacing = 1.0
                p_img.paragraph_format.space_before = Pt(0)
                p_img.paragraph_format.space_after = Pt(0)
                
                photo_path = resolve_photo_path(photo.get("file_path", ""))
                compat_path, is_temp = get_docx_compatible_image(photo_path)
                if compat_path:
                    try:
                        p_img.add_run().add_picture(compat_path, width=Cm(6))
                    except Exception as e:
                        print(f"Exception adding picture: {e}")
                        run_pl = p_img.add_run(f"[Nota: {photo.get('caption', 'Belanja')}]")
                        format_run(run_pl, size_pt=11, italic=True)
                    finally:
                        if is_temp and os.path.exists(compat_path):
                            try:
                                os.remove(compat_path)
                            except:
                                pass
                else:
                    run_pl = p_img.add_run("[Nota Belanja]")
                    format_run(run_pl, size_pt=11, italic=True)
                    
                cap_cell = table_nota.rows[grid_row + 1].cells[grid_col]
                set_cell_margins(cap_cell, top=40, bottom=80, left=100, right=100)
                p_cap = cap_cell.paragraphs[0]
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.line_spacing = 1.0
                p_cap.paragraph_format.space_before = Pt(0)
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(photo.get("caption", ""))
                format_run(run_cap, size_pt=11, italic=True)
        
    if is_mubesma:
        # D. PROGRAM KERJA YANG TIDAK TEREALISASI
        p_hdr_d = doc.add_paragraph()
        p_hdr_d.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_d.paragraph_format.space_before = Pt(12)
        p_hdr_d.paragraph_format.space_after = Pt(6)
        p_hdr_d.paragraph_format.keep_with_next = True
        pref_d_title = f"{pref_d} PROGRAM KERJA YANG TIDAK TEREALISASI" if is_mubesma else f"{pref_d} PROGRAM KERJA YANG BELUM TEREALISASI"
        format_run(p_hdr_d.add_run(pref_d_title), size_pt=12, bold=True)
        
        pbt_list = config_data.get("proker_belum_terlaksana", [])
        if not pbt_list:
            p_pbt_empty = doc.add_paragraph()
            p_pbt_empty.paragraph_format.left_indent = Cm(0.5)
            format_run(p_pbt_empty.add_run("(Tidak ada program kerja belum terlaksana)"), size_pt=11, italic=True)
        
        for idx, pk in enumerate(pbt_list, 1):
            p_name = doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(12)
            p_name.paragraph_format.space_after = Pt(6)
            p_name.paragraph_format.keep_with_next = True
            p_name.paragraph_format.left_indent = Cm(0.5)
            
            name_str = clean_proker_name(pk.get('Nama Kegiatan', pk.get('Nama Program Kerja', 'Proker')))
            format_run(p_name.add_run(f"{idx}. {name_str}"), size_pt=12, bold=True)
            
            table = doc.add_table(rows=9, cols=3)
            fields = [
                ("Nama Kegiatan", clean_proker_name(pk.get("Nama Kegiatan", pk.get("Nama Program Kerja", "")))),
                ("Sifat", pk.get("Sifat", "—")),
                ("Tema Kegiatan", pk.get("Tema Kegiatan", "—")),
                ("Tujuan Kegiatan", pk.get("Tujuan Kegiatan", pk.get("Tujuan", "—"))),
                ("Tanggal Kegiatan", format_id_date_str(pk.get("Tanggal Kegiatan", "—"))),
                ("Penanggung Jawab", pk.get("Penanggung Jawab", "—")),
                ("Peserta Kegiatan", pk.get("Peserta Kegiatan", "—")),
                ("Anggaran", pk.get("Anggaran", "—")),
                ("Dokumentasi", pk.get("Dokumentasi", "—"))
            ]
            render_table_rows(table, fields, indent_cm=INDENT_PROKER, bold_label=False, prefix_alpha=True)
            
    if is_mubesma:
        # E. EVALUASI KINERJA MENTERI
        p_hdr_e = doc.add_paragraph()
        p_hdr_e.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_e.paragraph_format.space_before = Pt(12)
        p_hdr_e.paragraph_format.space_after = Pt(6)
        p_hdr_e.paragraph_format.keep_with_next = True
        format_run(p_hdr_e.add_run(f"{pref_e} EVALUASI KINERJA MENTERI"), size_pt=12, bold=True)

        eval_pribadi = config_data.get("evaluasi_kinerja_pribadi", "").strip()
        if eval_pribadi:
            for line in eval_pribadi.split("\n"):
                line_clean = line.strip()
                if not line_clean:
                    doc.add_paragraph()
                    continue
                p_eval_pribadi = doc.add_paragraph()
                p_eval_pribadi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_eval_pribadi.paragraph_format.line_spacing = 1.5
                p_eval_pribadi.paragraph_format.space_before = Pt(0)
                p_eval_pribadi.paragraph_format.space_after = Pt(0)
                p_eval_pribadi.paragraph_format.left_indent = Cm(0.5)
                p_eval_pribadi.paragraph_format.first_line_indent = Cm(1.0)
                format_run(p_eval_pribadi.add_run(line_clean), size_pt=12)
        else:
            p_eval_pribadi = doc.add_paragraph()
            p_eval_pribadi.paragraph_format.left_indent = Cm(0.5)
            p_eval_pribadi.paragraph_format.space_before = Pt(0)
            p_eval_pribadi.paragraph_format.space_after = Pt(0)
            p_eval_pribadi.paragraph_format.line_spacing = 1.5
            format_run(p_eval_pribadi.add_run("—"), size_pt=12)

        # PENUTUP SECTION
        p_hdr_penutup = doc.add_paragraph()
        p_hdr_penutup.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_hdr_penutup.paragraph_format.space_before = Pt(12)
        p_hdr_penutup.paragraph_format.space_after = Pt(6)
        p_hdr_penutup.paragraph_format.keep_with_next = True
        format_run(p_hdr_penutup.add_run(f"{pref_penutup} PENUTUP"), size_pt=12, bold=True)

        penutup_text = config_data.get("penutup", "").strip()
        if not penutup_text:
            penutup_text = "Demikian Laporan Pertanggungjawaban ini kami susun sebagai bentuk pertanggungjawaban atas amanah yang telah diberikan selama satu periode kepengurusan. Kami menyadari masih banyak kekurangan dalam pelaksanaan program kerja maupun dalam koordinasi internal, namun hal tersebut menjadi bahan evaluasi dan pembelajaran untuk ke depannya.\n\nTerima kasih kepada seluruh pihak yang telah mendukung dan bekerja sama, baik dari internal maupun pihak eksternal. Semoga apa yang telah dijalankan dapat memberikan manfaat bagi mahasiswa dan lingkungan kampus secara luas."

        for line in penutup_text.split("\n"):
            line_clean = line.strip()
            if not line_clean:
                doc.add_paragraph()
                continue
            p_penutup = doc.add_paragraph()
            p_penutup.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_penutup.paragraph_format.line_spacing = 1.5
            p_penutup.paragraph_format.space_before = Pt(0)
            p_penutup.paragraph_format.space_after = Pt(0)
            p_penutup.paragraph_format.left_indent = Cm(0.5)
            p_penutup.paragraph_format.first_line_indent = Cm(1.0)
            format_run(p_penutup.add_run(line_clean), size_pt=12)

    # Signature Block
    import re
    clean_k_name = re.sub(r'^(Kementerian|Menteri|Departemen)\s+', '', kementrian_str, flags=re.IGNORECASE)
    if not clean_k_name or clean_k_name.lower() == 'kementerian':
        clean_k_name = 'Luar Kampus'
    org_prefix = "Menteri" if "bem" in output_path.lower() else "Menteri"
    
    import datetime
    months = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    now = datetime.datetime.now()
    tgl_str = f"Majalengka, {now.day:02d} {months[now.month - 1]} {now.year}"
    
    if is_mubesma:
        p_sig_date = doc.add_paragraph()
        p_sig_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig_date.paragraph_format.space_before = Pt(36)
        p_sig_date.paragraph_format.line_spacing = 1.15
        format_run(p_sig_date.add_run(tgl_str), size_pt=12)

        p_sig_ket1 = doc.add_paragraph()
        p_sig_ket1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig_ket1.paragraph_format.line_spacing = 1.15
        format_run(p_sig_ket1.add_run("Ketua Umum"), size_pt=12)

        p_sig_ket2 = doc.add_paragraph()
        p_sig_ket2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig_ket2.paragraph_format.line_spacing = 1.15
        format_run(p_sig_ket2.add_run(f"{org_prefix} {clean_k_name},\n\n\n\n\n"), size_pt=12)

        ketua_name = config_data.get("keanggotaan", {}).get("ketua", "Nama Ketua")

        p_sig_name = doc.add_paragraph()
        p_sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig_name.paragraph_format.line_spacing = 1.15
        run_name = p_sig_name.add_run(ketua_name)
        run_name.underline = True
        format_run(run_name, size_pt=12, bold=True)

    add_document_footer(doc, triwulan_str)
    doc.save(output_path)
    return True

def consolidate_lpj(output_path, file_list):
    master_doc = docx.Document()
    set_document_formatting(master_doc)
    
    # 1. Master Cover
    p_cover = master_doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cover.paragraph_format.line_spacing = 1.15
    p_cover.paragraph_format.space_before = Pt(72)
    p_cover.paragraph_format.space_after = Pt(24)
    
    triwulan_val = "X"
    periode_val = "2025-2026"
    
    parsed_docs_data = []
    
    for f in file_list:
        pdata = parse_docx(f)
        if pdata:
            parsed_docs_data.append((f, pdata))
            if pdata["cover"]["triwulan"]:
                triwulan_val = pdata["cover"]["triwulan"]
            if pdata["cover"]["periode"]:
                periode_val = pdata["cover"]["periode"]
                
    if triwulan_val.upper() == "MUBESMA":
        run_mc1 = p_cover.add_run("LAPORAN PERTANGGUNGJAWABAN MUBESMA\n\n")
    else:
        run_mc1 = p_cover.add_run(f"LAPORAN PERTANGGUNGJAWABAN TRIWULAN {triwulan_val.upper()}\n\n")
    format_run(run_mc1, size_pt=14, bold=True)
    
    run_mc2 = p_cover.add_run("BADAN EKSEKUTIF MAHASISWA (BEM)\n")
    format_run(run_mc2, size_pt=14, bold=True)
    
    run_mc3 = p_cover.add_run(f"INSTBUNAS MAJALENGKA\nPERIODE {periode_val}\n\n\n\n\n\n\n\n")
    format_run(run_mc3, size_pt=14, bold=True)
    
    master_doc.add_page_break()
    
    # 2. Table of Contents Placeholder
    p_toc_title = master_doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.space_after = Pt(12)
    format_run(p_toc_title.add_run("DAFTAR ISI\n"), size_pt=14, bold=True)
    
    p_toc = master_doc.add_paragraph()
    p_toc.paragraph_format.line_spacing = 1.5
    
    is_mubesma = (triwulan_val.upper() == "MUBESMA")
    for idx, (f, pdata) in enumerate(parsed_docs_data):
        kementerian_name = pdata["cover"]["kementerian"] or os.path.basename(f)
        ch_letter = int_to_roman(idx + 1) if is_mubesma else chr(65 + idx)
        p_toc.add_run(f"{ch_letter}. LAPORAN MENTERI {kementerian_name.upper()} ").font.name = "Times New Roman"
        p_toc.add_run("." * 60 + "\n").font.name = "Times New Roman"
        
    final_letter = int_to_roman(len(parsed_docs_data) + 1) if is_mubesma else chr(65 + len(parsed_docs_data))
    p_toc.add_run(f"{final_letter}. RINGKASAN ANGGARAN TERPADU BEM ").font.name = "Times New Roman"
    p_toc.add_run("." * 60 + "\n").font.name = "Times New Roman"
    
    master_doc.add_page_break()
    
    # 3. Append Kementerian documents
    for idx, (f, pdata) in enumerate(parsed_docs_data):
        ch_letter = int_to_roman(idx + 1) if is_mubesma else chr(65 + idx)
        k_name = pdata["cover"]["kementerian"] or "Kementerian"
        
        if triwulan_val.upper() == "MUBESMA":
            add_minister_cover(master_doc, k_name, pdata["cover"]["periode"] or periode_val)
            
        p_ch = master_doc.add_paragraph()
        p_ch.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ch.paragraph_format.space_before = Pt(12)
        p_ch.paragraph_format.space_after = Pt(6)
        p_ch.paragraph_format.keep_with_next = True
        format_run(p_ch.add_run(f"{ch_letter}. LAPORAN PERTANGGUNGJAWABAN MENTERI {k_name.upper()}"), size_pt=12, bold=True)
        
        # Keadaan Objektif / Pendahuluan
        p_sub1 = master_doc.add_paragraph()
        p_sub1.paragraph_format.space_before = Pt(12)
        p_sub1.paragraph_format.space_after = Pt(6)
        p_sub1.paragraph_format.keep_with_next = True
        sub1_title = "1. Pendahuluan" if is_mubesma else "1. Keadaan Objektif Menteri"
        format_run(p_sub1.add_run(sub1_title), size_pt=12, bold=True)
        
        obj_text = pdata["keadaan_objektif"].strip() if pdata["keadaan_objektif"] else "Deskripsi Keadaan Objektif..."
        for line in obj_text.split("\n"):
            line_clean = line.strip()
            if not line_clean:
                master_doc.add_paragraph()
                continue
            p_obj = master_doc.add_paragraph()
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_obj.paragraph_format.line_spacing = 1.5
            p_obj.paragraph_format.space_before = Pt(0)
            p_obj.paragraph_format.space_after = Pt(0)
            p_obj.paragraph_format.left_indent = Cm(0.5)
            p_obj.paragraph_format.first_line_indent = Cm(1.0)
            format_run(p_obj.add_run(line_clean), size_pt=12)
        
        # Keanggotaan
        p_sub2 = master_doc.add_paragraph()
        p_sub2.paragraph_format.space_before = Pt(12)
        p_sub2.paragraph_format.space_after = Pt(6)
        p_sub2.paragraph_format.keep_with_next = True
        format_run(p_sub2.add_run("2. Susunan Keanggotaan"), size_pt=12, bold=True)
        
        anggota_val = pdata['keanggotaan'].get('anggota', '')
        has_anggota = False
        if anggota_val:
            if isinstance(anggota_val, list):
                anggota_str = ", ".join(anggota_val).strip()
            else:
                anggota_str = str(anggota_val).strip()
            if anggota_str and anggota_str != "-" and anggota_str.lower() != "null":
                has_anggota = True
                
        rows_data = [
            ("Ketua Menteri", pdata['keanggotaan'].get('ketua', '')),
            ("Sekretaris", pdata['keanggotaan'].get('sekretaris', '')),
            ("Bendahara", pdata['keanggotaan'].get('bendahara', ''))
        ]
        if has_anggota:
            rows_data.append(("Anggota", anggota_str))
            
        table_k = master_doc.add_table(rows=len(rows_data), cols=3)
        render_table_rows(table_k, rows_data)
        
        if is_mubesma:
            p_sub_tup = master_doc.add_paragraph()
            p_sub_tup.paragraph_format.space_before = Pt(12)
            p_sub_tup.paragraph_format.space_after = Pt(6)
            p_sub_tup.paragraph_format.keep_with_next = True
            format_run(p_sub_tup.add_run("3. Tugas Pokok dan Fungsi"), size_pt=12, bold=True)
            
            # A. Tugas Pokok
            p_tp = master_doc.add_paragraph()
            p_tp.paragraph_format.space_before = Pt(6)
            p_tp.paragraph_format.space_after = Pt(4)
            p_tp.paragraph_format.keep_with_next = True
            p_tp.paragraph_format.left_indent = Cm(0.5)
            format_run(p_tp.add_run("A. Tugas Pokok"), size_pt=12, bold=True)
            
            tugas_list = pdata.get("tugas_pokok", [])
            if not tugas_list:
                tugas_list = ["(Belum diatur)"]
            for t_idx, t_item in enumerate(tugas_list, 1):
                p_t_item = master_doc.add_paragraph()
                p_t_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_t_item.paragraph_format.line_spacing = 1.15
                p_t_item.paragraph_format.space_after = Pt(4)
                p_t_item.paragraph_format.left_indent = Cm(1.5)
                p_t_item.paragraph_format.first_line_indent = Cm(-0.5)
                format_run(p_t_item.add_run(f"{t_idx}.\t{t_item}"), size_pt=12)
                
            # B. Fungsi
            p_f = master_doc.add_paragraph()
            p_f.paragraph_format.space_before = Pt(6)
            p_f.paragraph_format.space_after = Pt(4)
            p_f.paragraph_format.keep_with_next = True
            p_f.paragraph_format.left_indent = Cm(0.5)
            format_run(p_f.add_run("B. Fungsi"), size_pt=12, bold=True)
            
            fungsi_list = pdata.get("fungsi", [])
            if not fungsi_list:
                fungsi_list = ["(Belum diatur)"]
            for f_idx, f_item in enumerate(fungsi_list, 1):
                p_f_item = master_doc.add_paragraph()
                p_f_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_f_item.paragraph_format.line_spacing = 1.15
                p_f_item.paragraph_format.space_after = Pt(4)
                p_f_item.paragraph_format.left_indent = Cm(1.5)
                p_f_item.paragraph_format.first_line_indent = Cm(-0.5)
                format_run(p_f_item.add_run(f"{f_idx}.\t{f_item}"), size_pt=12)
                
            # 4. Evaluasi Pencapaian Visi dan Misi
            p_sub_vm = master_doc.add_paragraph()
            p_sub_vm.paragraph_format.space_before = Pt(12)
            p_sub_vm.paragraph_format.space_after = Pt(6)
            p_sub_vm.paragraph_format.keep_with_next = True
            format_run(p_sub_vm.add_run("4. Evaluasi Pencapaian Visi dan Misi"), size_pt=12, bold=True)
            
            p_sub_vm_lbl = master_doc.add_paragraph()
            p_sub_vm_lbl.paragraph_format.space_before = Pt(6)
            p_sub_vm_lbl.paragraph_format.space_after = Pt(6)
            p_sub_vm_lbl.paragraph_format.keep_with_next = True
            p_sub_vm_lbl.paragraph_format.left_indent = Cm(0.5)
            org_lbl = "Badan Perwakilan Mahasiswa" if "bpm" in output_path.lower() else "Badan Eksekutif Mahasiswa"
            format_run(p_sub_vm_lbl.add_run(f"Visi dan Misi {org_lbl} INSTBUNAS Majalengka"), size_pt=12, bold=True)
            
            # A. Visi
            p_v_hdr = master_doc.add_paragraph()
            p_v_hdr.paragraph_format.space_before = Pt(6)
            p_v_hdr.paragraph_format.space_after = Pt(4)
            p_v_hdr.paragraph_format.keep_with_next = True
            p_v_hdr.paragraph_format.left_indent = Cm(0.5)
            format_run(p_v_hdr.add_run("A. Visi"), size_pt=12, bold=True)
            
            visi_text = pdata.get("visi", "").strip()
            if not visi_text:
                visi_text = "(Belum diatur)"
            p_v_txt = master_doc.add_paragraph()
            p_v_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_v_txt.paragraph_format.line_spacing = 1.15
            p_v_txt.paragraph_format.space_after = Pt(12)
            p_v_txt.paragraph_format.left_indent = Cm(0.5)
            format_run(p_v_txt.add_run(visi_text), size_pt=12)
            
            # B. Misi
            p_m_hdr = master_doc.add_paragraph()
            p_m_hdr.paragraph_format.space_before = Pt(6)
            p_m_hdr.paragraph_format.space_after = Pt(4)
            p_m_hdr.paragraph_format.keep_with_next = True
            p_m_hdr.paragraph_format.left_indent = Cm(0.5)
            format_run(p_m_hdr.add_run("B. Misi"), size_pt=12, bold=True)
            
            misi_list = pdata.get("misi", [])
            if not misi_list:
                misi_list = ["(Belum diatur)"]
            for m_idx, m_item in enumerate(misi_list, 1):
                p_m_item = master_doc.add_paragraph()
                p_m_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_m_item.paragraph_format.line_spacing = 1.15
                p_m_item.paragraph_format.space_after = Pt(4)
                p_m_item.paragraph_format.left_indent = Cm(1.5)
                p_m_item.paragraph_format.first_line_indent = Cm(-0.5)
                format_run(p_m_item.add_run(f"{m_idx}.\t{m_item}"), size_pt=12)
        
        # Sort proker terlaksana chronologically
        proker_terlaksana = pdata.get("proker_terlaksana", [])
        proker_terlaksana.sort(key=lambda x: parse_id_date(x.get('Tanggal Kegiatan', '')))

        if is_mubesma:
            # 5. Program Kerja (Summary Table)
            p_sub_sum = master_doc.add_paragraph()
            p_sub_sum.paragraph_format.space_before = Pt(12)
            p_sub_sum.paragraph_format.space_after = Pt(6)
            p_sub_sum.paragraph_format.keep_with_next = True
            format_run(p_sub_sum.add_run("5. Program Kerja"), size_pt=12, bold=True)
            
            # Create summary table
            table_v = master_doc.add_table(rows=1, cols=5)
            table_v.style = 'Table Grid'
            set_table_indent(table_v, 0.5)
            set_table_widths(table_v, [Cm(1.0), Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.0)])
            
            headers_v = ["NO", "WAKTU", "PROGRAM KERJA", "NAMA KEGIATAN", "TEMPAT"]
            hdr_cells = table_v.rows[0].cells
            for c_idx, h in enumerate(headers_v):
                hdr_cells[c_idx].text = h
                set_cell_shading(hdr_cells[c_idx], "D3D3D3")
                set_cell_margins(hdr_cells[c_idx], top=100, bottom=100, left=120, right=120)
                hdr_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[c_idx].paragraphs[0].paragraph_format.line_spacing = 1.0
                format_run(hdr_cells[c_idx].paragraphs[0].runs[0], size_pt=12, bold=True)
                
            proker_groups_v = []
            prev_prog = None
            current_no = 1
            for pk_row in proker_terlaksana:
                prog_name = clean_proker_name(pk_row.get('Nama Program Kerja', '—'))
                if prev_prog == prog_name and proker_groups_v:
                    proker_groups_v[-1]['rows'].append(pk_row)
                else:
                    proker_groups_v.append({
                        'name': pk_row.get('Nama Program Kerja', '—'),
                        'start_no': current_no,
                        'rows': [pk_row]
                    })
                    prev_prog = prog_name
                current_no += 1
                
            row_idx = 1
            for grp in proker_groups_v:
                span = len(grp['rows'])
                start_row = row_idx
                for r_idx, r_row in enumerate(grp['rows']):
                    row_cells = table_v.add_row().cells
                    set_table_widths(table_v, [Cm(1.0), Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.0)])
                    
                    if r_idx == 0:
                        row_cells[0].text = str(grp['start_no'])
                        row_cells[2].text = grp['name']
                    
                    row_cells[1].text = format_id_date_str(r_row.get('Tanggal Kegiatan', '—'))
                    row_cells[3].text = r_row.get('Nama Kegiatan', r_row.get('Nama Program Kerja', '—'))
                    row_cells[4].text = r_row.get('Tempat Kegiatan', r_row.get('Tempat', '—'))
                    
                    for c_idx, cell in enumerate(row_cells):
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        cell.vertical_alignment = 0
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for r in p.runs:
                            format_run(r, size_pt=12)
                    row_idx += 1
                    
                end_row = row_idx - 1
                if span > 1:
                    cell_no_start = table_v.cell(start_row, 0)
                    cell_no_end = table_v.cell(end_row, 0)
                    cell_no_start.merge(cell_no_end)
                    cell_no_start.vertical_alignment = 0
                    
                    cell_pk_start = table_v.cell(start_row, 2)
                    cell_pk_end = table_v.cell(end_row, 2)
                    cell_pk_start.merge(cell_pk_end)
                    cell_pk_start.vertical_alignment = 0

        # Realisasi Proker
        p_sub3 = master_doc.add_paragraph()
        p_sub3.paragraph_format.space_before = Pt(12)
        p_sub3.paragraph_format.space_after = Pt(6)
        p_sub3.paragraph_format.keep_with_next = True
        sub3_title = "6. Program Kerja Yang Terealisasi" if is_mubesma else "3. Program Kerja Yang Terealisasi"
        format_run(p_sub3.add_run(sub3_title), size_pt=12, bold=True)
        
        for idx_pk, pk in enumerate(proker_terlaksana, 1):
            p_name = master_doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(12)
            p_name.paragraph_format.space_after = Pt(6)
            p_name.paragraph_format.keep_with_next = True
            
            name_str = clean_proker_name(pk.get('Nama Program Kerja', 'Proker'))
            format_run(p_name.add_run(f"{idx_pk}. {name_str}"), size_pt=12, bold=True)
            
            table = master_doc.add_table(rows=9, cols=3)
            fields = [
                ("Nama Kegiatan", clean_proker_name(pk.get("Nama Kegiatan", pk.get("Nama Program Kerja", "")))),
                ("Tempat Kegiatan", pk.get("Tempat Kegiatan", pk.get("Tempat", ""))),
                ("Sifat", pk.get("Sifat", "Internal")),
                ("Tema Kegiatan", pk.get("Tema Kegiatan", "")),
                ("Tujuan", pk.get("Tujuan", "")),
                ("Tanggal Kegiatan", format_id_date_str(pk.get("Tanggal Kegiatan", ""))),
                ("Penanggung Jawab", pk.get("Penanggung Jawab", "")),
                ("Peserta Kegiatan", pk.get("Peserta Kegiatan", "")),
                ("Evaluasi & Saran", pk.get("Evaluasi & Saran", pk.get("Evaluasi", "")))
            ]
            INDENT_PROKER = 1.0
            render_table_rows(table, fields, indent_cm=INDENT_PROKER, bold_label=False, prefix_alpha=True)
            
            # Sub-bagian: Realisasi Anggaran
            p_sub_ang = master_doc.add_paragraph()
            p_sub_ang.paragraph_format.space_before = Pt(12)
            p_sub_ang.paragraph_format.space_after = Pt(6)
            p_sub_ang.paragraph_format.keep_with_next = True
            p_sub_ang.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_sub_ang.add_run("j.  Realisasi Anggaran"), size_pt=12, bold=False)
            
            tidak_menggunakan_anggaran = pk.get("tidak_menggunakan_anggaran", False)
            anggaran_list = pk.get("anggaran", [])
            
            if tidak_menggunakan_anggaran or not anggaran_list:
                p_ang_empty = master_doc.add_paragraph()
                p_ang_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
                format_run(p_ang_empty.add_run("(Tidak ada realisasi anggaran)"), size_pt=12, italic=True)
            else:
                table_ang = master_doc.add_table(rows=1, cols=6)
                set_table_indent(table_ang, INDENT_PROKER)
                table_ang.style = 'Table Grid'
                set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
                
                headers = ["Tanggal", "Keterangan", "Uraian", "Debet", "Kredit", "Saldo"]
                hdr_cells = table_ang.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = h
                    set_cell_shading(hdr_cells[i], "D3D3D3")
                    set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hdr_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
                    format_run(hdr_cells[i].paragraphs[0].runs[0], size_pt=12, bold=True)
                    
                running_balance = 0
                total_debet = 0
                total_kredit = 0
                
                for tx in anggaran_list:
                    row_cells = table_ang.add_row().cells
                    set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
                    
                    deb_val = clean_currency(tx.get("debet", 0))
                    kred_val = clean_currency(tx.get("kredit", 0))
                    
                    total_debet += deb_val
                    total_kredit += kred_val
                    
                    if running_balance == 0 and len(table_ang.rows) == 2:
                        running_balance = deb_val - kred_val
                    else:
                        running_balance = running_balance + deb_val - kred_val
                        
                    row_cells[0].text = tx.get("tanggal", "")
                    row_cells[1].text = tx.get("keterangan", "")
                    row_cells[2].text = tx.get("uraian", "")
                    row_cells[3].text = format_currency(deb_val) if deb_val > 0 else ""
                    row_cells[4].text = format_currency(kred_val) if kred_val > 0 else ""
                    row_cells[5].text = format_currency(running_balance)
                    
                    for idx_c, cell in enumerate(row_cells):
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        if idx_c in [3, 4, 5]:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for r in p.runs:
                            format_run(r, size_pt=12)
                            
                tot_cells = table_ang.add_row().cells
                set_table_widths(table_ang, [Cm(1.5), Cm(2.8), Cm(2.6), Cm(2.0), Cm(2.0), Cm(2.1)])
                
                tot_cells[0].text = "TOTAL"
                tot_cells[1].text = "TOTAL"
                tot_cells[2].text = "TOTAL"
                tot_cells[3].text = format_currency(total_debet)
                tot_cells[4].text = format_currency(total_kredit)
                tot_cells[5].text = format_currency(running_balance)
                
                for idx_c, cell in enumerate(tot_cells):
                    set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                    set_cell_shading(cell, "EAEAEA")
                    p = cell.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if idx_c in [3, 4, 5]:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        format_run(r, size_pt=12, bold=True)
                        
            # Sub-bagian: Dokumentasi Kegiatan
            p_sub_dok = master_doc.add_paragraph()
            p_sub_dok.paragraph_format.space_before = Pt(12)
            p_sub_dok.paragraph_format.space_after = Pt(6)
            p_sub_dok.paragraph_format.keep_with_next = True
            p_sub_dok.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_sub_dok.add_run("k.  Dokumentasi Kegiatan"), size_pt=12, bold=False)
            
            doc_list = pk.get("dokumentasi", [])
            num_photos = len(doc_list)
            
            if num_photos == 0:
                p_dok_empty = master_doc.add_paragraph()
                p_dok_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
                format_run(p_dok_empty.add_run("(Dokumentasi tidak tersedia)"), size_pt=12, italic=True)
            else:
                num_rows_doc = ((num_photos + 1) // 2) * 2
                table_doc = master_doc.add_table(rows=num_rows_doc, cols=2)
                set_table_indent(table_doc, INDENT_PROKER)
                remove_table_borders(table_doc)
                set_table_widths(table_doc, [Cm(6.5), Cm(6.5)])
                
                for i, photo in enumerate(doc_list):
                    grid_row = (i // 2) * 2
                    grid_col = i % 2
                    
                    img_cell = table_doc.rows[grid_row].cells[grid_col]
                    set_cell_margins(img_cell, top=80, bottom=40, left=100, right=100)
                    p_img = img_cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.line_spacing = 1.0
                    p_img.paragraph_format.space_before = Pt(0)
                    p_img.paragraph_format.space_after = Pt(0)
                    
                    photo_path = resolve_photo_path(photo.get("file_path", ""))
                    compat_path, is_temp = get_docx_compatible_image(photo_path)
                    if compat_path:
                        try:
                            p_img.add_run().add_picture(compat_path, width=Cm(6))
                        except Exception as e:
                            print(f"Exception adding picture: {e}")
                            run_pl = p_img.add_run(f"[Foto: {photo.get('caption', 'Kegiatan')}]")
                            format_run(run_pl, size_pt=11, italic=True)
                        finally:
                            if is_temp and os.path.exists(compat_path):
                                try:
                                    os.remove(compat_path)
                                except:
                                    pass
                    else:
                        run_pl = p_img.add_run("[Foto Kegiatan]")
                        format_run(run_pl, size_pt=11, italic=True)
                        
                    cap_cell = table_doc.rows[grid_row + 1].cells[grid_col]
                    set_cell_margins(cap_cell, top=40, bottom=80, left=100, right=100)
                    p_cap = cap_cell.paragraphs[0]
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.line_spacing = 1.0
                    p_cap.paragraph_format.space_before = Pt(0)
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(photo.get("caption", ""))
                    format_run(run_cap, size_pt=11, italic=True)

            # Sub-bagian: Nota Belanja
            p_sub_nota = master_doc.add_paragraph()
            p_sub_nota.paragraph_format.space_before = Pt(12)
            p_sub_nota.paragraph_format.space_after = Pt(6)
            p_sub_nota.paragraph_format.keep_with_next = True
            p_sub_nota.paragraph_format.left_indent = Cm(INDENT_PROKER)
            format_run(p_sub_nota.add_run("l.  Nota Belanja"), size_pt=12, bold=False)
            
            nota_list = pk.get("nota_belanja", [])
            num_notas = len(nota_list)
            
            if num_notas == 0:
                p_nota_empty = master_doc.add_paragraph()
                p_nota_empty.paragraph_format.left_indent = Cm(INDENT_PROKER)
                format_run(p_nota_empty.add_run("(Nota belanja tidak tersedia)"), size_pt=12, italic=True)
            else:
                num_rows_nota = ((num_notas + 1) // 2) * 2
                table_nota = master_doc.add_table(rows=num_rows_nota, cols=2)
                set_table_indent(table_nota, INDENT_PROKER)
                remove_table_borders(table_nota)
                set_table_widths(table_nota, [Cm(6.5), Cm(6.5)])
                
                for i, photo in enumerate(nota_list):
                    grid_row = (i // 2) * 2
                    grid_col = i % 2
                    
                    img_cell = table_nota.rows[grid_row].cells[grid_col]
                    set_cell_margins(img_cell, top=80, bottom=40, left=100, right=100)
                    p_img = img_cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.line_spacing = 1.0
                    p_img.paragraph_format.space_before = Pt(0)
                    p_img.paragraph_format.space_after = Pt(0)
                    
                    photo_path = resolve_photo_path(photo.get("file_path", ""))
                    compat_path, is_temp = get_docx_compatible_image(photo_path)
                    if compat_path:
                        try:
                            p_img.add_run().add_picture(compat_path, width=Cm(6))
                        except Exception as e:
                            print(f"Exception adding picture: {e}")
                            run_pl = p_img.add_run(f"[Nota: {photo.get('caption', 'Belanja')}]")
                            format_run(run_pl, size_pt=11, italic=True)
                        finally:
                            if is_temp and os.path.exists(compat_path):
                                try:
                                    os.remove(compat_path)
                                except:
                                    pass
                    else:
                        run_pl = p_img.add_run("[Nota Belanja]")
                        format_run(run_pl, size_pt=11, italic=True)
                        
                    cap_cell = table_nota.rows[grid_row + 1].cells[grid_col]
                    set_cell_margins(cap_cell, top=40, bottom=80, left=100, right=100)
                    p_cap = cap_cell.paragraphs[0]
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.line_spacing = 1.0
                    p_cap.paragraph_format.space_before = Pt(0)
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(photo.get("caption", ""))
                    format_run(run_cap, size_pt=11, italic=True)
            
        # Program Kerja Belum Realisasi
        p_sub4 = master_doc.add_paragraph()
        p_sub4.paragraph_format.space_before = Pt(12)
        p_sub4.paragraph_format.space_after = Pt(6)
        p_sub4.paragraph_format.keep_with_next = True
        sub4_title = "7. Program Kerja Tidak Terealisasi" if is_mubesma else "4. Program Kerja Belum Terealisasi"
        format_run(p_sub4.add_run(sub4_title), size_pt=12, bold=True)
        
        for idx_pk, pk in enumerate(pdata["proker_belum_terlaksana"], 1):
            p_name = master_doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(12)
            p_name.paragraph_format.space_after = Pt(6)
            p_name.paragraph_format.keep_with_next = True
            
            name_str = pk.get('Nama Kegiatan', pk.get('Nama Program Kerja', 'Proker'))
            format_run(p_name.add_run(f"{idx_pk}. {name_str}"), size_pt=12, bold=True)
            
            table = master_doc.add_table(rows=9, cols=3)
            fields = [
                ("Nama Kegiatan", pk.get("Nama Kegiatan", pk.get("Nama Program Kerja", ""))),
                ("Sifat", pk.get("Sifat", "—")),
                ("Tema Kegiatan", pk.get("Tema Kegiatan", "—")),
                ("Tujuan Kegiatan", pk.get("Tujuan Kegiatan", pk.get("Tujuan", "—"))),
                ("Tanggal Kegiatan", format_id_date_str(pk.get("Tanggal Kegiatan", "—"))),
                ("Penanggung Jawab", pk.get("Penanggung Jawab", "—")),
                ("Peserta Kegiatan", pk.get("Peserta Kegiatan", "—")),
                ("Anggaran", pk.get("Anggaran", "—")),
                ("Dokumentasi", pk.get("Dokumentasi", "—"))
            ]
            render_table_rows(table, fields)
            
        # Evaluasi Kinerja Menteri
        p_sub5 = master_doc.add_paragraph()
        p_sub5.paragraph_format.space_before = Pt(12)
        p_sub5.paragraph_format.space_after = Pt(6)
        p_sub5.paragraph_format.keep_with_next = True
        sub5_title = "8. Evaluasi Kinerja Menteri" if is_mubesma else "5. Evaluasi Kinerja Pribadi"
        format_run(p_sub5.add_run(sub5_title), size_pt=12, bold=True)
        
        val_ep = pdata.get("evaluasi_kinerja_pribadi", "").strip()
        if val_ep:
            for line in val_ep.split("\n"):
                line_clean = line.strip()
                if not line_clean:
                    master_doc.add_paragraph()
                    continue
                p_eval_prib = master_doc.add_paragraph()
                p_eval_prib.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_eval_prib.paragraph_format.line_spacing = 1.5
                p_eval_prib.paragraph_format.space_before = Pt(0)
                p_eval_prib.paragraph_format.space_after = Pt(0)
                p_eval_prib.paragraph_format.left_indent = Cm(0.5)
                p_eval_prib.paragraph_format.first_line_indent = Cm(1.0)
                format_run(p_eval_prib.add_run(line_clean), size_pt=12)
        else:
            p_eval_prib = master_doc.add_paragraph()
            p_eval_prib.paragraph_format.left_indent = Cm(0.5)
            p_eval_prib.paragraph_format.space_before = Pt(0)
            p_eval_prib.paragraph_format.space_after = Pt(0)
            p_eval_prib.paragraph_format.line_spacing = 1.5
            format_run(p_eval_prib.add_run("—"), size_pt=12)
        
        if not is_mubesma:
            # Evaluasi Anggota dan Internal Menteri
            p_sub6 = master_doc.add_paragraph()
            p_sub6.paragraph_format.space_before = Pt(12)
            p_sub6.paragraph_format.space_after = Pt(6)
            p_sub6.paragraph_format.keep_with_next = True
            sub6_title = "6. Evaluasi Anggota dan Internal Menteri"
            format_run(p_sub6.add_run(sub6_title), size_pt=12, bold=True)
            
            eval_anggota = pdata.get("evaluasi_anggota_internal", [])
            if eval_anggota:
                table_eva = master_doc.add_table(rows=1, cols=4)
                table_eva.style = 'Table Grid'
                set_table_indent(table_eva, 0.5)
                set_table_widths(table_eva, [Cm(1.0), Cm(4.0), Cm(5.5), Cm(5.5)])
                
                hdr_cells = table_eva.rows[0].cells
                hdr_cells[0].text = 'No'
                hdr_cells[1].text = 'Nama Anggota'
                hdr_cells[2].text = 'Kepribadian'
                hdr_cells[3].text = 'Kinerja'
                
                for cell in hdr_cells:
                    set_cell_shading(cell, "D3D3D3")
                    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        format_run(run, size_pt=11, bold=True)
                        
                for idx_eva, agt in enumerate(eval_anggota, 1):
                    row_cells = table_eva.add_row().cells
                    set_table_widths(table_eva, [Cm(1.0), Cm(4.0), Cm(5.5), Cm(5.5)])
                    row_cells[0].text = str(idx_eva)
                    row_cells[1].text = agt.get('nama', '')
                    row_cells[2].text = agt.get('kepribadian', '')
                    row_cells[3].text = agt.get('kinerja', '')
                    
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for i_cell, cell in enumerate(row_cells):
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.line_spacing = 1.15
                            paragraph.paragraph_format.space_after = Pt(2)
                            for run in paragraph.runs:
                                format_run(run, size_pt=10 if i_cell > 1 else 11)
            else:
                p_eval_agt = master_doc.add_paragraph()
                p_eval_agt.paragraph_format.left_indent = Cm(0.5)
                run_ea = p_eval_agt.add_run("—")
                format_run(run_ea, size_pt=11)
            
        master_doc.add_page_break()
        
    # 4. Ringkasan Anggaran Terpadu BEM
    letter_akhir = int_to_roman(len(parsed_docs_data) + 1) if is_mubesma else chr(65 + len(parsed_docs_data))
    p_final_ch = master_doc.add_paragraph()
    p_final_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_final_ch.paragraph_format.space_before = Pt(12)
    p_final_ch.paragraph_format.space_after = Pt(12)
    p_final_ch.paragraph_format.keep_with_next = True
    format_run(p_final_ch.add_run(f"{letter_akhir}. RINGKASAN ANGGARAN TERPADU BEM INSTBUNAS"), size_pt=12, bold=True)
    
    table_summary = master_doc.add_table(rows=1, cols=4)
    table_summary.style = 'Table Grid'
    
    headers_sum = ["Kementerian", "Total Debet", "Total Kredit", "Saldo"]
    hdr_sum_cells = table_summary.rows[0].cells
    for c_idx, h in enumerate(headers_sum):
        hdr_sum_cells[c_idx].text = h
        set_cell_shading(hdr_sum_cells[c_idx], "D3D3D3")
        set_cell_margins(hdr_sum_cells[c_idx], top=100, bottom=100, left=120, right=120)
        hdr_sum_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_sum_cells[c_idx].paragraphs[0].paragraph_format.line_spacing = 1.0
        format_run(hdr_sum_cells[c_idx].paragraphs[0].runs[0], size_pt=12, bold=True)
        
    grand_debet = 0
    grand_kredit = 0
    grand_saldo = 0
    
    for f, pdata in parsed_docs_data:
        k_name = pdata["cover"]["kementerian"] or os.path.basename(f)
        deb = pdata["anggaran_summary"]["debet"]
        kred = pdata["anggaran_summary"]["kredit"]
        sal = pdata["anggaran_summary"]["saldo"]
        
        grand_debet += deb
        grand_kredit += kred
        grand_saldo += sal
        
        row_cells = table_summary.add_row().cells
        row_cells[0].text = k_name
        row_cells[1].text = format_currency(deb)
        row_cells[2].text = format_currency(kred)
        row_cells[3].text = format_currency(sal)
        
        for c_idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if c_idx > 0: 
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                format_run(r, size_pt=12)
                
    gt_cells = table_summary.add_row().cells
    gt_cells[0].text = "GRAND TOTAL"
    gt_cells[1].text = format_currency(grand_debet)
    gt_cells[2].text = format_currency(grand_kredit)
    gt_cells[3].text = format_currency(grand_saldo)
    
    for c_idx, cell in enumerate(gt_cells):
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        set_cell_shading(cell, "EAEAEA")
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if c_idx > 0: 
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            format_run(r, size_pt=12, bold=True)
            
    add_document_footer(master_doc, triwulan_val)
    master_doc.save(output_path)
    return True

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python3 scratch/bem_lpj_manager.py validate <file_path>")
        print("  python3 scratch/bem_lpj_manager.py generate <output_path> <input_json_path>")
        print("  python3 scratch/bem_lpj_manager.py consolidate <output_path> <file1> <file2> ...")
        sys.exit(1)
        
    action = sys.argv[1].lower()
    
    if action == "validate":
        file_path = sys.argv[2]
        res = run_validation(file_path)
        if len(sys.argv) > 3 and sys.argv[3] == "--json":
            print(json.dumps(res))
            sys.exit(0)
        if res.get("status") == "ERROR":
            print(res["message"])
            sys.exit(1)
            
        print("CHECKLIST VALIDASI LPJ KEMENTERIAN")
        print("────────────────────────────────────────────────────────────────")
        print(f"{'✅' if res['checklist']['cover'] else '❌'}  Cover: Nomor Triwulan, Nama Kementerian, Periode tercantum")
        print(f"{'✅' if res['checklist']['keanggotaan'] else '❌'}  Keanggotaan: Ketua, Sekretaris, Bendahara terisi semua")
        print(f"{'✅' if res['checklist']['proker_terlaksana'] else '❌'}  Minimal 1 Proker terlaksana terdokumentasi lengkap (8 field)")
        print(f"{'✅' if res['checklist']['tabel_anggaran_kolom'] and res['checklist']['tabel_anggaran_saldo'] else '❌'}  Tabel anggaran: kolom lengkap, saldo berjalan konsisten")
        print(f"{'✅' if res['checklist']['tabel_anggaran_total'] else '❌'}  Baris TOTAL anggaran ada dan kalkulasi benar")
        print(f"{'✅' if res['checklist']['proker_belum_terlaksana'] else '❌'}  Proker belum terlaksana: tercantum + ada target tanggal")
        print(f"{'✅' if res['checklist']['dokumentasi'] else '❌'}  Dokumentasi: minimal ada keterangan foto (walau foto kosong)")
        print("────────────────────────────────────────────────────────────────")
        print(f"STATUS: {res['status']}")
        
        if res['errors']:
            print("\nCATATAN:")
            for idx, err in enumerate(res['errors'], 1):
                print(f"{idx}. {err}")
        else:
            print("\nCatatan: Semua field dan perhitungan data valid.")
            
    elif action == "generate":
        out_path = sys.argv[2]
        json_path = sys.argv[3]
        with open(json_path, 'r') as f:
            cdata = json.load(f)
        success = generate_lpj(out_path, cdata)
        if success:
            print(f"LPJ successfully generated at: {out_path}")
        else:
            print("Failed to generate LPJ.")
            
    elif action == "consolidate":
        out_path = sys.argv[2]
        files = sys.argv[3:]
        success = consolidate_lpj(out_path, files)
        if success:
            print(f"Consolidated LPJ successfully created at: {out_path}")
        else:
            print("Failed to consolidate LPJ files.")
